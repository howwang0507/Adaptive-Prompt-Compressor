import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            # Simple bold/italic handling
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(docx_path)
    print(f"Success: {docx_path} generated.")


if __name__ == "__main__":
    md_file = "paper_draft.md"
    docx_file = "Research_Paper_Adaptive_Prompt_Compression.docx"
    markdown_to_docx(md_file, docx_file)
